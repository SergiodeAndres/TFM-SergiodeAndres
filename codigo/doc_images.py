import os
from tkinter import Tk, filedialog
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

TITULOS = [
    "Tabla de temas de semestres de 2016 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2017 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2018 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2019 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2020 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2021 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2022 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2023 de Elon Musk (GenSim - LDA)",
    "Tabla de temas de semestres de 2024 de Elon Musk (GenSim - LDA)"
]

TEXTO_MARCADOR = "[AQUÍ EMPIEZAN LAS IMÁGENES]"

# --- FUNCIONES ---
def insertar_parrafo_despues(parrafo, texto="", estilo=None):
    new_p = OxmlElement("w:p")
    parrafo._element.addnext(new_p)
    nuevo_parrafo = parrafo._parent.add_paragraph(texto, style=estilo)
    return nuevo_parrafo

# --- SELECCIÓN DE ARCHIVOS ---
Tk().withdraw()
doc_path = filedialog.askopenfilename(title="Selecciona el archivo de Word", filetypes=[("Word files", "*.docx")])
if not doc_path:
    raise Exception("No se seleccionó ningún archivo.")
doc = Document(doc_path)

img_paths = filedialog.askopenfilenames(title="Selecciona las imágenes", filetypes=[("Imágenes", "*.png;*.jpg;*.jpeg;*.bmp")])
if not img_paths:
    raise Exception("No se seleccionaron imágenes.")

if len(img_paths) > len(TITULOS):
    raise Exception("Hay más imágenes que títulos. Ajusta la lista TITULOS.")

# --- BUSCAR EL MARCADOR ---
marcador_parrafo = None
for para in doc.paragraphs:
    if TEXTO_MARCADOR in para.text:
        marcador_parrafo = para
        break

if not marcador_parrafo:
    raise Exception(f"No se encontró el marcador: {TEXTO_MARCADOR}")

# --- INSERTAR IMÁGENES Y TÍTULOS ---
for i, img_path in enumerate(img_paths):
    titulo = TITULOS[i]

    # Insertar imagen
    p_img = insertar_parrafo_despues(marcador_parrafo)
    run = p_img.add_run()
    run.add_picture(img_path, width=Cm(19.5))
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.left_indent = Cm(-2.2)

    # Insertar título
    p_title = insertar_parrafo_despues(p_img, texto=titulo, estilo='Caption')
    p_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Avanzar el marcador
    marcador_parrafo = p_title

# --- GUARDAR DOCUMENTO ---
nuevo_path = os.path.splitext(doc_path)[0] + "_con_imagenes.docx"
doc.save(nuevo_path)
print(f"✅ Documento guardado como: {nuevo_path}")